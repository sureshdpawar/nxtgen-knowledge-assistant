from pathlib import Path

from docx import (
    Document,
)

from app.parsers.base import (
    BaseParser,
    ParsedResult,
)


class DocxParser(
    BaseParser
):

    def extract(
        self,
        file_path: Path,
    ) -> ParsedResult:

        document = (
            Document(
                str(
                    file_path
                )
            )
        )

        text_parts: list[str] = []

        for paragraph in (
            document.paragraphs
        ):
            text = (
                paragraph.text
                .strip()
            )

            if text:
                text_parts.append(
                    text
                )

        for table in (
            document.tables
        ):
            for row in table.rows:
                values = [
                    cell.text.strip()
                    for cell
                    in row.cells
                ]

                if any(
                    values
                ):
                    text_parts.append(
                        " | ".join(
                            values
                        )
                    )

        text = (
            "\n".join(
                text_parts
            )
            .strip()
        )

        return {
            "pages": [
                {
                    "page": 1,
                    "text": text,
                }
            ],

            "metadata": {
                "page_count": 1,
                "document_type": "word",
            },
        }